import google.generativeai as genai
import os
import re
import time
import threading
import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk, filedialog
from tkinter.font import Font
from dotenv import load_dotenv
import yt_dlp
from pathlib import Path
import logging
import sys

# Configure logging with encoding support
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log", encoding="utf-8"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class VideoDownloaderApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Video Quiz Generator")
        self.root.geometry("900x750")
        self.root.minsize(800, 650)
        
        # Set theme
        style = ttk.Style()
        if 'clam' in style.theme_names():
            style.theme_use('clam')
        
        # Custom styles
        style.configure('TLabelframe', borderwidth=1)
        style.configure('TLabelframe.Label', font=('Segoe UI', 10, 'bold'))
        style.configure('TButton', font=('Segoe UI', 9))
        style.configure('Primary.TButton', background='#4285F4', foreground='white')
        
        # Set up variables
        self.video_url = tk.StringVar()
        self.num_questions = tk.StringVar(value="5")  # Default value
        self.output_dir = tk.StringVar(value=os.path.join(os.getcwd(), "downloads"))
        self.is_processing = False
        
        # Ensure output directory exists
        Path(self.output_dir.get()).mkdir(parents=True, exist_ok=True)
        
        # Main frame
        main_frame = ttk.Frame(root, padding="15 15 15 15")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # App header
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        # App title
        title_font = Font(family="Segoe UI", size=16, weight="bold")
        app_title = ttk.Label(header_frame, text="YouTube Video Quiz Generator", font=title_font)
        app_title.pack(side=tk.LEFT)
        
        # Create UI elements
        self.create_ui(main_frame)
        
        # Load API key after UI is created
        self.load_api_key()
    
    def load_api_key(self):
        """Load environment variables and configure Gemini API"""
        load_dotenv()
        api_key = os.getenv('API_KEY')
        
        if not api_key:
            self.log_update("API_KEY not found in environment variables!", error=True)
            messagebox.showerror("API Key Missing", 
                                "Please create a .env file with your API_KEY or set it as an environment variable.")
        else:
            try:
                genai.configure(api_key=api_key)
                self.log_update("Gemini API configured successfully")
            except Exception as e:
                self.log_update(f"Failed to configure Gemini API: {e}", error=True)
    
    def create_ui(self, parent):
        """Create the application UI"""
        # Create frame for inputs with a more modern design
        input_frame = ttk.LabelFrame(parent, text="Input Settings", padding="10 10 10 10")
        input_frame.pack(fill="x", expand=False, padx=0, pady=(0, 10))
        
        # Grid configuration for input frame
        input_frame.columnconfigure(1, weight=1)
        
        # URL input with improved layout
        url_label = ttk.Label(input_frame, text="YouTube URL:")
        url_label.grid(row=0, column=0, padx=5, pady=10, sticky="w")
        
        url_entry = ttk.Entry(input_frame, textvariable=self.video_url, width=70)
        url_entry.grid(row=0, column=1, padx=5, pady=10, sticky="ew")
        url_entry.focus()  # Set focus to URL entry on startup
        
        # Questions input with better spacing
        questions_label = ttk.Label(input_frame, text="Number of Questions:")
        questions_label.grid(row=1, column=0, padx=5, pady=10, sticky="w")
        
        questions_frame = ttk.Frame(input_frame)
        questions_frame.grid(row=1, column=1, padx=5, pady=10, sticky="w")
        
        questions_spinbox = ttk.Spinbox(questions_frame, from_=1, to=20, textvariable=self.num_questions, width=5)
        questions_spinbox.pack(side=tk.LEFT)
        
        hint_label = ttk.Label(questions_frame, text="(1-20 questions)", font=("Segoe UI", 9, "italic"), foreground="gray")
        hint_label.pack(side=tk.LEFT, padx=10)
        
        # Output directory with improved design
        dir_label = ttk.Label(input_frame, text="Save Location:")
        dir_label.grid(row=2, column=0, padx=5, pady=10, sticky="w")
        
        dir_frame = ttk.Frame(input_frame)
        dir_frame.grid(row=2, column=1, padx=5, pady=10, sticky="ew")
        
        dir_entry = ttk.Entry(dir_frame, textvariable=self.output_dir)
        dir_entry.pack(side=tk.LEFT, fill="x", expand=True)
        
        browse_button = ttk.Button(dir_frame, text="Browse", command=self.browse_directory)
        browse_button.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Status section
        status_frame = ttk.Frame(parent)
        status_frame.pack(fill="x", padx=0, pady=(0, 10))
        
        # Status label with icon
        self.status_var = tk.StringVar(value="Ready")
        self.status_label = ttk.Label(status_frame, textvariable=self.status_var, font=("Segoe UI", 9))
        self.status_label.pack(side=tk.LEFT, anchor="w", padx=(5, 0))
        
        # Progress bar with improved appearance
        self.progress = ttk.Progressbar(parent, orient="horizontal", mode="determinate", style="TProgressbar")
        self.progress.pack(fill="x", padx=0, pady=(0, 10))
        
        # Buttons frame with better layout
        button_frame = ttk.Frame(parent)
        button_frame.pack(fill="x", pady=(0, 15))
        
        # Start button (primary)
        self.start_button = ttk.Button(
            button_frame, 
            text="Start Analysis", 
            command=self.start_analysis,
            style="Primary.TButton",
            padding=(20, 5)
        )
        self.start_button.pack(side=tk.LEFT, padx=(0, 10))
        
        # Cancel button
        self.cancel_button = ttk.Button(
            button_frame, 
            text="Cancel", 
            command=self.cancel_operation, 
            state=tk.DISABLED,
            padding=(20, 5)
        )
        self.cancel_button.pack(side=tk.LEFT)
        
        # Clear log button to the right
        self.clear_button = ttk.Button(
            button_frame, 
            text="Clear Log", 
            command=self.clear_log,
            padding=(10, 5)
        )
        self.clear_button.pack(side=tk.RIGHT)
        
        # Activity log with improved appearance
        log_frame = ttk.LabelFrame(parent, text="Activity Log", padding="10 10 10 10")
        log_frame.pack(fill="both", expand=True, padx=0, pady=0)
        
        # More attractive log box with better font
        self.log_box = scrolledtext.ScrolledText(
            log_frame, 
            wrap=tk.WORD, 
            font=("Consolas", 10),
            background="#F8F8F8",
            borderwidth=1,
            relief="solid"
        )
        self.log_box.pack(fill="both", expand=True)
        
        # Tag configuration for different message types
        self.log_box.tag_config("error", foreground="#D32F2F")  # Red
        self.log_box.tag_config("success", foreground="#388E3C")  # Green
        self.log_box.tag_config("info", foreground="#1976D2")  # Blue
        
        # Status bar at the bottom
        status_bar = ttk.Frame(parent, relief=tk.SUNKEN, padding=(5, 2))
        status_bar.pack(fill=tk.X, side=tk.BOTTOM)
        
        version_label = ttk.Label(status_bar, text="v1.0.0", font=("Segoe UI", 8))
        version_label.pack(side=tk.RIGHT)
        
        # Initial log message
        self.log_update("Application ready. Enter a YouTube URL and click 'Start Analysis'.")
    
    def browse_directory(self):
        """Open file dialog to select output directory"""
        directory = filedialog.askdirectory(initialdir=self.output_dir.get())
        if directory:
            self.output_dir.set(directory)
            self.log_update(f"Output directory set to: {directory}")
    
    def log_update(self, message, error=False, success=False):
        """Update the log box in the GUI with formatted messages"""
        timestamp = time.strftime("%H:%M:%S")
        formatted_message = f"[{timestamp}] {message}"
        
        # Log to file as well
        if error:
            logger.error(message)
            tag = "error"
        elif success:
            logger.info(message)
            tag = "success"
        else:
            logger.info(message)
            tag = "info"
        
        # Insert into log box with appropriate tags
        self.log_box.insert(tk.END, formatted_message + '\n', tag)
        self.log_box.yview(tk.END)  # Scroll to the bottom
        self.log_box.update_idletasks()  # Ensure the GUI updates immediately
    
    def clear_log(self):
        """Clear the log box"""
        self.log_box.delete(1.0, tk.END)
        self.log_update("Log cleared")
    
    def update_progress(self, percentage=None, status=None):
        """Update progress bar and status label"""
        if percentage is not None:
            self.progress["value"] = percentage
        
        if status:
            self.status_var.set(status)
            
        self.root.update_idletasks()
    
    def extract_youtube_id(self, youtube_url):
        """Extract the YouTube video ID from a given URL"""
        patterns = [
            r"(?:v=|\/)([A-Za-z0-9_-]{11})(?:\?|&|\/|$)",  # Standard and embed URLs
            r"(?:youtu\.be\/)([A-Za-z0-9_-]{11})(?:\?|&|$)",  # Short URLs
            r"(?:shorts\/)([A-Za-z0-9_-]{11})(?:\?|&|$)"  # YouTube Shorts
        ]
        
        for pattern in patterns:
            match = re.search(pattern, youtube_url)
            if match:
                return match.group(1)
        
        return None
    
    def download_video(self, video_url):
        """Download video and update progress"""
        try:
            # Generate a unique filename based on video ID
            video_id = self.extract_youtube_id(video_url)
            if not video_id:
                self.log_update("Invalid YouTube URL format", error=True)
                return None, None
            
            output_filename = os.path.join(self.output_dir.get(), f"{video_id}.mp4")
            
            # Skip download if the file already exists
            if os.path.exists(output_filename):
                self.log_update(f"Video file already exists: {output_filename}")
                response = messagebox.askyesno(
                    "File Exists", 
                    "Video has already been downloaded. Use existing file?\n\nClick 'No' to redownload."
                )
                if response:
                    # Get video duration
                    with yt_dlp.YoutubeDL({'quiet': True}) as ydl:
                        info_dict = ydl.extract_info(video_url, download=False)
                        self.log_update(f"Using existing video: {info_dict.get('title')}")
                        return output_filename, info_dict.get('duration')
                else:
                    self.log_update(f"Deleting existing video file for redownload")
                    os.remove(output_filename)
            
            # Set up progress hooks for download tracking
            def progress_hook(d):
                if d['status'] == 'downloading':
                    if 'total_bytes' in d and 'downloaded_bytes' in d:
                        percentage = (d['downloaded_bytes'] / d['total_bytes']) * 100
                        self.update_progress(
                            percentage=percentage, 
                            status=f"Downloading: {percentage:.1f}% - ETA: {d.get('eta', '?')}s"
                        )
                    
                    # Update log every 5% or when ETA changes significantly
                    if d.get('_percent_str', ''):
                        self.log_update(f"Downloading {d.get('_percent_str', '')} of {d.get('_total_bytes_str', '?')} "
                                       f"at {d.get('_speed_str', '?')} ETA {d.get('_eta_str', '?')}")
                
                elif d['status'] == 'finished':
                    self.log_update(f"Download complete! Processing video...", success=True)
                    self.update_progress(percentage=100, status="Download complete")
            
            ydl_opts = {
                'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best',
                'outtmpl': output_filename,
                'progress_hooks': [progress_hook],
                'quiet': True,  # Suppress output to console
                'no_warnings': True
            }
            
            self.log_update(f"Starting download of YouTube video: {video_url}")
            self.update_progress(percentage=0, status="Starting download...")
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info_dict = ydl.extract_info(video_url, download=True)
                filename = ydl.prepare_filename(info_dict)
                self.log_update(f"Video downloaded: {info_dict.get('title')}", success=True)
                self.log_update(f"File saved as: {filename}")
                
                return filename, info_dict.get('duration')
                
        except Exception as e:
            self.log_update(f"Error downloading video: {str(e)}", error=True)
            logger.exception("Download error")
            return None, None
    
    def upload_video_to_gemini(self, video_file_path):
        """Upload video to Gemini and ensure it's ready for inference"""
        try:
            self.log_update(f"Uploading video file to Gemini...")
            self.update_progress(status="Uploading to Gemini...")
            
            video_file = genai.upload_file(video_file_path)
            self.log_update(f"Upload initiated with URI: {video_file.uri}")
            
            # Monitor processing status
            attempts = 0
            max_attempts = 20  # Maximum number of status checks
            
            while video_file.state.name != "ACTIVE" and attempts < max_attempts:
                attempts += 1
                wait_time = min(10, attempts)  # Gradually increase wait time
                
                self.update_progress(
                    status=f"Processing video... (Check {attempts}/{max_attempts})")
                self.log_update(f"Video processing: {video_file.state.name} - Waiting {wait_time}s...")
                
                time.sleep(wait_time)
                video_file = genai.get_file(video_file.name)
            
            if video_file.state.name == "FAILED":
                self.log_update("File processing failed", error=True)
                return None
            elif video_file.state.name != "ACTIVE":
                self.log_update("File processing timed out", error=True)
                return None
            
            self.log_update("File processed and ready for analysis", success=True)
            return video_file
            
        except Exception as e:
            self.log_update(f"Error uploading video to Gemini: {str(e)}", error=True)
            logger.exception("Upload error")
            return None
    
    def analyze_video_with_gemini(self, video_file, num_questions):
        """Analyze video with Gemini and return the response"""
        try:
            self.log_update(f"Starting video analysis with Gemini...")
            self.update_progress(status="Analyzing video content...")
            
            model = genai.GenerativeModel(model_name="gemini-1.5-pro")
            
            # Construct a detailed, instructive prompt
            prompt = (
                f"Create a comprehensive quiz from this video with exactly {num_questions} multiple choice questions.\n\n"
                "Requirements:\n"
                "1. Start with a concise summary of the video's main points (2-3 paragraphs)\n"
                f"2. Create exactly {num_questions} multiple choice questions (A, B, C, D format)\n"
                "3. Include the exact timestamp (HH:MM:SS format) where each question's content appears\n"
                "4. Distribute questions evenly throughout the video (beginning, middle, and end)\n"
                "5. Only reference information explicitly stated in the video\n"
                "6. Present questions in chronological order\n"
                "7. Provide an answer key at the end with brief explanations\n\n"
                "Format the output as a well-structured markdown document with clear headings."
            )
            
            self.log_update("Sending inference request to Gemini (this may take several minutes)...")
            
            # Set a longer timeout for the API request
            response = model.generate_content(
                [video_file, prompt], 
                request_options={"timeout": 600}
            )
            
            if not response or not response.text:
                self.log_update("Empty response from Gemini API", error=True)
                return None
                
            self.log_update("Analysis completed successfully!", success=True)
            return response.text
            
        except Exception as e:
            self.log_update(f"Error during video analysis: {str(e)}", error=True)
            logger.exception("Analysis error")
            return None
    
    def save_quiz_to_file(self, quiz_content, video_id):
        """Save the quiz content to a file"""
        try:
            # Create a base filename from the video ID
            base_filename = f"quiz_{video_id}_{time.strftime('%Y%m%d_%H%M%S')}"
            
            # Ask user for file format preference with improved dialog
            file_choice = messagebox.askyesnocancel(
                "Save Quiz", 
                "Save quiz results?\n\nYes: Save as .txt\nNo: Save as .docx\nCancel: Don't save",
                icon='question'
            )
            
            if file_choice is None:  # User canceled
                self.log_update("Quiz not saved (user canceled)")
                return
                
            if file_choice:  # Save as TXT
                filepath = os.path.join(self.output_dir.get(), f"{base_filename}.txt")
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(quiz_content)
                self.log_update(f"Quiz saved as: {filepath}", success=True)
                
            else:  # Save as DOCX
                try:
                    from docx import Document
                    filepath = os.path.join(self.output_dir.get(), f"{base_filename}.docx")
                    doc = Document()
                    
                    # Split content by lines and paragraphs for better formatting
                    lines = quiz_content.split('\n')
                    current_paragraph = []
                    
                    for line in lines:
                        # Check if line is a heading (starts with #)
                        if line.strip().startswith('#'):
                            # Write any accumulated paragraph text
                            if current_paragraph:
                                doc.add_paragraph(''.join(current_paragraph))
                                current_paragraph = []
                                
                            # Add heading with appropriate level
                            level = min(line.count('#'), 9)  # Word supports up to 9 heading levels
                            text = line.strip('#').strip()
                            doc.add_heading(text, level=level)
                        
                        # Check if line is empty (paragraph break)
                        elif not line.strip():
                            if current_paragraph:
                                doc.add_paragraph(''.join(current_paragraph))
                                current_paragraph = []
                        
                        # Regular content line
                        else:
                            current_paragraph.append(line + " ")
                    
                    # Add any remaining paragraph
                    if current_paragraph:
                        doc.add_paragraph(''.join(current_paragraph))
                    
                    doc.save(filepath)
                    self.log_update(f"Quiz saved as: {filepath}", success=True)
                    
                except ImportError:
                    self.log_update("python-docx package not installed. Saving as TXT instead.", error=True)
                    # Fallback to TXT
                    filepath = os.path.join(self.output_dir.get(), f"{base_filename}.txt")
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(quiz_content)
                    self.log_update(f"Quiz saved as: {filepath}", success=True)
            
            # Ask if user wants to open the file
            if messagebox.askyesno("File Saved", "Would you like to open the saved file?"):
                # Handle opening files in a platform-independent way
                if os.name == 'nt':  # Windows
                    os.startfile(filepath)
                elif os.name == 'posix':  # macOS and Linux
                    os.system(f"open '{filepath}'" if sys.platform == 'darwin' else f"xdg-open '{filepath}'")
                
        except Exception as e:
            self.log_update(f"Error saving quiz: {str(e)}", error=True)
            logger.exception("Save error")
    
    def cancel_operation(self):
        """Cancel the current operation"""
        if self.is_processing:
            self.is_processing = False
            self.log_update("Operation canceled by user", error=True)
            self.update_progress(status="Operation canceled")
            self.toggle_processing_state(False)
    
    def toggle_processing_state(self, is_processing):
        """Update UI elements based on processing state"""
        self.is_processing = is_processing
        
        if is_processing:
            self.start_button.config(state=tk.DISABLED)
            self.cancel_button.config(state=tk.NORMAL)
        else:
            self.start_button.config(state=tk.NORMAL)
            self.cancel_button.config(state=tk.DISABLED)
            self.progress["value"] = 0
            self.update_progress(status="Ready")
    
    def start_analysis(self):
        """Start the video download and analysis process"""
        # Validate input
        video_url = self.video_url.get().strip()
        if not video_url:
            self.log_update("Please enter a valid YouTube URL", error=True)
            messagebox.showerror("Input Error", "Please enter a valid YouTube URL")
            return
            
        try:
            num_questions = int(self.num_questions.get())
            if num_questions < 1 or num_questions > 20:
                raise ValueError("Number of questions must be between 1 and 20")
        except ValueError as e:
            self.log_update(f"Invalid number of questions: {str(e)}", error=True)
            messagebox.showerror("Input Error", f"Invalid number of questions: {str(e)}")
            return
            
        # Start processing thread
        self.toggle_processing_state(True)
        threading.Thread(target=self.process_video, args=(video_url, num_questions), daemon=True).start()
    
    def process_video(self, video_url, num_questions):
        """Process the video in a separate thread"""
        try:
            video_id = self.extract_youtube_id(video_url)
            if not video_id:
                self.log_update("Could not extract video ID from URL", error=True)
                self.toggle_processing_state(False)
                return
                
            # Step 1: Download video
            video_file_path, video_duration = self.download_video(video_url)
            
            if not video_file_path or not video_duration:
                self.log_update("Video download failed", error=True)
                self.toggle_processing_state(False)
                return
                
            # Check video duration limit (Gemini has a 30-minute limit)
            if video_duration > 1800:  # 30 minutes in seconds
                self.log_update(f"Video duration ({video_duration/60:.1f} minutes) exceeds Gemini's 30-minute limit", error=True)
                messagebox.showwarning(
                    "Video Too Long", 
                    f"The video is {video_duration/60:.1f} minutes long, which exceeds Gemini's 30-minute limit.\n\n"
                    "The analysis may be incomplete or fail. Do you want to continue anyway?",
                )
            
            # Step 2: Upload to Gemini
            video_file = self.upload_video_to_gemini(video_file_path)
            if not video_file:
                self.log_update("Video upload failed", error=True)
                self.toggle_processing_state(False)
                return
                
            # Step 3: Analyze video
            analysis_result = self.analyze_video_with_gemini(video_file, num_questions)
            if not analysis_result:
                self.log_update("Video analysis failed", error=True)
                self.toggle_processing_state(False)
                return
                
            # Step 4: Save results
            self.save_quiz_to_file(analysis_result, video_id)
            
            # Complete
            self.log_update("Video analysis completed successfully!", success=True)
            self.update_progress(status="Analysis complete")
            messagebox.showinfo("Success", "Video analysis completed successfully!")
            
        except Exception as e:
            self.log_update(f"Error during processing: {str(e)}", error=True)
            logger.exception("Processing error")
            messagebox.showerror("Error", f"An error occurred during processing: {str(e)}")
            
        finally:
            # Always reset the UI state
            self.toggle_processing_state(False)

def main():
    """Main application entry point"""
    try:
        root = tk.Tk()
        # Set app icon if available
        try:
            root.iconbitmap("app_icon.ico")  # You'll need to create this icon
        except:
            pass  # Ignore if icon not found
            
        app = VideoDownloaderApp(root)
        root.mainloop()
    except Exception as e:
        logger.exception(f"Application crashed: {e}")
        # Handle errors that occur before GUI is available
        try:
            messagebox.showerror("Fatal Error", f"The application encountered a fatal error: {str(e)}")
        except:
            print(f"FATAL ERROR: {str(e)}")

if __name__ == "__main__":
    main()